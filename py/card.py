# -*- coding: utf-8 -*-
"""개별 응답 카드

한 사람이 한 게임에 대해 답한 내용을 전부 담는다.
평가 문장은 넣지 않는다. 답과 그 답이 전체 분포에서 어디에 놓이는지만 적는다.

  표지        게임명 · 응답자 ID · 유형 · 제출 일시
  한눈에      본인 6축 평균 · 최고/최저 항목 · 추천 점수 · 진행도 · 구매 의향 · 적정가
  1. 응답자   연령 · 성별 · 선호 장르 · 소속 집단과 그 집단의 평균
  2. 평가 점수 육각 차트(본인 vs 전체) · 8항목 표(내 점수 · 평균 · 차이 · 순위)
  3. 게임 문항 문항별로 내 답과 전체 분포를 함께
  4. 공통 문항 진행도 · 구매 의향 · 적정가 · 버그
  5. 서술형    문항 전문과 답 그대로. 무응답도 적는다
  6. 관련 이슈 이 응답자의 답이 근거로 들어간 종합 리포트 3장 항목

응답자 실명은 어디에도 쓰지 않는다.
"""
import os

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from . import config as C
from . import diagrams as D
from .aggregate import 백분율, 찾기
from .docxkit import (ACCENT, GRAY, SEV, body_p, card, caption, cell_margins,
                      datatable, 문서만들기, fixed_table, h, pic, run, shade,
                      spacer, white_borders)
from .ingest import 무응답, 숫자
from .kor import 은는, 이가, 자르기


def _KPI(doc, r, 지표):
    """본인 기준 요약 6칸"""
    본인점수 = [v for v in r.육각 if v is not None]
    본인평균 = sum(본인점수) / len(본인점수) if 본인점수 else None
    쌍 = [(v, 이름) for v, (_, 이름) in zip(r.육각, C.육각축) if v is not None]
    최고 = max(쌍, default=None)
    최저 = min(쌍, default=None)
    구분 = ""
    if r.추천의향 is not None:
        구분 = "추천" if r.추천의향 >= 9 else ("중립" if r.추천의향 >= 7 else "비추천")

    항목 = [
        ("6축 평균", f"{본인평균:.2f}" if 본인평균 is not None else "-",
         f'전체 {지표["육각총평균"]:.2f}' if 지표["육각총평균"] is not None else ""),
        ("가장 높게", 최고[1] if 최고 else "-", f"{최고[0]}점" if 최고 else "-"),
        ("가장 낮게", 최저[1] if 최저 else "-", f"{최저[0]}점" if 최저 else "-"),
        ("추천 점수", f"{r.추천의향}" if r.추천의향 is not None else "-", 구분),
        ("진행도", 자르기(r.진행도 or "-", 12), ""),
        ("적정가", 자르기(r.적정가격 or "-", 12), ""),
    ]
    t = fixed_table(doc, [1606] * 6)
    for i, (라벨, 값, 부) in enumerate(항목):
        c = t.rows[0].cells[i]
        white_borders(c); shade(c, "F1F6FB"); cell_margins(c, 120, 90, 120, 90)
        c.paragraphs[0].text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        run(c.paragraphs[0], 라벨, size=8, color=GRAY, bold=True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        run(p2, 값, size=10.5 if len(값) > 5 else 13, bold=True, color=ACCENT)
        p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        run(p3, 부, size=7.5, color=GRAY)
    spacer(doc, 10)


def _순위번호(값, 전체값들):
    """전체에서 몇 번째로 높은 점수인가 (같은 값은 공동 순위)"""
    유효 = sorted([v for v in 전체값들 if v is not None], reverse=True)
    if 값 is None or not 유효 or 값 not in 유효:
        return None, len(유효)
    return 유효.index(값) + 1, len(유효)


def _순위(값, 전체값들):
    등수, 총 = _순위번호(값, 전체값들)
    return "-" if 등수 is None else f"{등수}위 / {총}명"


def _분포줄(doc, 항, N, 강조답=None, 최대=6):
    """전체 분포를 한 줄로. 본인이 고른 선택지는 굵게."""
    if not 항["집계"]:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run(p, "전체  ", size=8, color=GRAY, bold=True)
    보임 = 항["집계"][:최대]
    for i, row in enumerate(보임):
        내것 = 강조답 is not None and row["답"] in 강조답
        if i:
            run(p, "  ·  ", size=8, color="C9D3DE")
        run(p, f'{자르기(row["답"], 20)} {row["인원"]}({row["비율"]}%)',
            size=8.3, bold=내것, color=ACCENT if 내것 else "3A4A5A")
    남은 = len(항["집계"]) - len(보임)
    if 남은 > 0:
        run(p, f"  외 {남은}개", size=8, color=GRAY)


def _문항블록(doc, 항, r, N):
    """문항 하나 — 제목 / 내 답 / 전체 분포"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run(p, 항["문항"], size=9, bold=True)

    if 항["출처"] == "고유":
        셀 = next((c for c in r.고유 if c["문항"] == 항["문항"]), None)
        답 = 셀["답"] if 셀 else ""
        값들 = 셀["값"] if 셀 else []
    else:
        답 = getattr(r, 항["키"]) or ""
        값들 = [답] if 답 else []

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run(p, "응답  ", size=8, color=GRAY, bold=True)

    if not 답:
        run(p, "무응답", size=9.5, color=GRAY, italic=True)
    elif 항["유형"] == "선형배율" and 항["통계"]:
        st = 항["통계"]
        내점수 = 숫자(답)
        run(p, f"{답}점", size=10, bold=True, color=ACCENT)
        if 내점수 is not None and st["평균"] is not None:
            차 = 내점수 - st["평균"]
            run(p, f'   전체 평균 {st["평균"]:.2f} · 표준편차 '
                   f'{st["표준편차"]:.2f}' if st["표준편차"] is not None
                   else f'   전체 평균 {st["평균"]:.2f}', size=8.5, color=GRAY)
            run(p, f"   차이 {'+' if 차 > 0 else ''}{차:.2f}", size=8.5, color=GRAY)
    else:
        run(p, 답, size=9.5, bold=True, color=ACCENT)
        인원들 = [찾기(항["집계"], v)["인원"] for v in 값들]
        if 인원들:
            표시 = " · ".join(f"{n}/{N}" for n in 인원들)
            소수 = any(n and n / N < C.임계값["소수의견_비율"] for n in 인원들)
            run(p, f"   같은 답 {표시}" + ("  (소수 의견)" if 소수 else ""), size=8.5,
                color="B45309" if 소수 else GRAY)

    _분포줄(doc, 항, N, 강조답=값들)


def 만들기(r, 집계, 경로, 차트경로, 편차경로=None, 띠경로=None):
    지표, 게임 = 집계["지표"], 집계["게임"]
    N = 지표["N"]
    doc = 문서만들기()

    # ══ 표지 ══
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, C.사업명, size=8.5, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, f"{게임['게임명']} 개별 응답 카드", size=21, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
    제출 = r.타임스탬프 if isinstance(r.타임스탬프, str) else str(r.타임스탬프)
    run(p, f"응답자 {r.ID}  ·  {r.유형}  ·  제출 {제출}  ·  개발사 {게임['개발사']}  ·  "
           f"작성 {C.기관명}", size=9.5, color=GRAY)

    h(doc, "한눈에 보기", 13, ACCENT, before=4, after=6)
    _KPI(doc, r, 지표)

    # ══ 1. 응답자 ══
    h(doc, "1. 응답자", 14, ACCENT, before=8, after=5, rule=True)
    세그 = 집계["세그먼트"]
    코어장르 = 세그["코어장르"]
    코어냐 = any(g in r.선호장르 for g in 코어장르)
    소속 = 세그["장르"][0] if 코어냐 else 세그["장르"][1]
    datatable(doc, ["항목", "응답"], [
        ["연령대", r.연령대 or "-"],
        ["성별", r.성별 or "-"],
        ["선호 장르", " · ".join(r.선호장르) if r.선호장르 else "-"],
        ["응답자 유형", r.유형 or "-"],
        ["장르 집단", f'{소속["라벨"]} ({소속["인원"]}명)'],
        ["집단 평균", (f'6축 {소속["육각평균"]:.2f} · NPS {소속["NPS평균"]:.1f} · '
                       f'진행 흐름 {소속["진행흐름"]:.2f}')
                      if 소속["육각평균"] is not None else "-"],
    ], [1900, 7737], sizes=(8.5, 9))
    spacer(doc, 8)

    # ══ 2. 평가 점수 ══
    h(doc, "2. 평가 점수", 14, ACCENT, before=10, after=5, rule=True)
    if os.path.exists(차트경로):
        pic(doc, 차트경로, 10.8)
        caption(doc, f"육각형 6축 — 실선이 {r.ID}, 점선이 응답자 {N}명 평균")

    행 = []
    for i, (_, 이름) in enumerate(C.육각축):
        내점수, 축 = r.육각[i], 지표["육각"][i]
        차 = (내점수 - 축["평균"]) if (내점수 is not None and 축["평균"] is not None) else None
        행.append([이름, 내점수 if 내점수 is not None else "-",
                   f'{축["평균"]:.2f}' if 축["평균"] is not None else "-",
                   ("+" if 차 and 차 > 0 else "") + (f"{차:.2f}" if 차 is not None else "-"),
                   _순위(내점수, [x.육각[i] for x in 집계["응답"]])])
    for k, (키, 이름) in enumerate(C.추가축):
        내점수, 축 = getattr(r, 키), 지표["추가"][k]
        차 = (내점수 - 축["평균"]) if (내점수 is not None and 축["평균"] is not None) else None
        행.append([이름, 내점수 if 내점수 is not None else "-",
                   f'{축["평균"]:.2f}' if 축["평균"] is not None else "-",
                   ("+" if 차 and 차 > 0 else "") + (f"{차:.2f}" if 차 is not None else "-"),
                   _순위(내점수, [getattr(x, 키) for x in 집계["응답"]])])
    datatable(doc, ["항목", "내 점수", "전체 평균", "차이", "순위"], 행,
              [2900, 1500, 1700, 1500, 2037], sizes=(8.5, 8.8))
    spacer(doc, 6)

    if 편차경로:
        pic(doc, 편차경로, 13.4)
        caption(doc, "항목별 전체 평균 대비 편차 — 붉은 쪽이 평균보다 낮다")

    본인점수 = [v for v in r.육각 if v is not None]
    if 본인점수:
        평균 = sum(본인점수) / len(본인점수)
        전체 = 지표["육각총평균"]
        등수, 총 = _순위번호(round(평균, 4),
                             [round(sum(v for v in x.육각 if v is not None)
                                    / max(len([v for v in x.육각 if v is not None]), 1), 4)
                              for x in 집계["응답"]])
        차 = 평균 - 전체
        꼬리 = f" 응답자 {총}명 중 {등수}위다." if 등수 else ""
        body_p(doc, f"6축 평균은 {평균:.2f}점으로 전체 평균 {전체:.2f}점과 "
                    f"{abs(차):.2f}점 차이가 난다.{꼬리}", size=9, color=GRAY)

    # ══ 3. 게임 고유 문항 ══
    고유선택 = [항 for 항 in 집계["문항집계"] if 항["출처"] == "고유" and 항["유형"] != "서술형"]
    if 고유선택:
        doc.add_page_break()
        h(doc, f'3. {게임["게임명"]} 문항', 14, ACCENT, before=0, after=5, rule=True)
        body_p(doc, f"개발사가 따로 물은 문항이다. 문항마다 이 응답자의 답과 "
                    f"응답자 {N}명 전체의 분포를 함께 적었다.", size=8.5, color=GRAY, after=2)
        for 항 in 고유선택:
            _문항블록(doc, 항, r, N)
        spacer(doc, 8)

    # ══ 4. 공통 문항 ══
    공통선택 = [항 for 항 in 집계["문항집계"] if 항["출처"] == "공통"]
    if 공통선택:
        h(doc, "4. 공통 문항", 14, ACCENT, before=10, after=5, rule=True)
        if 띠경로:
            pic(doc, 띠경로, 14.6)
            caption(doc, "진행도 단계와 추천 점수에서의 위치")
        for 항 in 공통선택:
            _문항블록(doc, 항, r, N)
        spacer(doc, 8)

    # ══ 5. 서술형 ══
    doc.add_page_break()
    서술항목 = []
    for 항 in 집계["문항집계"]:
        if 항["출처"] == "고유" and 항["유형"] == "서술형":
            셀 = next((c for c in r.고유 if c["문항"] == 항["문항"]), None)
            서술항목.append((항["문항"], 셀["답"] if 셀 else "", len(항["서술"])))
    for 키, 라벨 in [("개선점", "현재 빌드에서 수정 · 개선되었으면 하는 점"),
                     ("버그상황", "버그 · 오류 발생 상황"),
                     ("개발팀메시지", "개발팀에 전하고 싶은 말")]:
        건수 = sum(1 for x in 집계["응답"] if not 무응답(getattr(x, 키)))
        서술항목.append((라벨, getattr(r, 키), 건수))

    쓴것 = sum(1 for _l, g, _c in 서술항목 if not 무응답(g))
    h(doc, f"5. 서술형 응답  ({쓴것}/{len(서술항목)}건 작성)", 14, ACCENT, before=0, after=5, rule=True)
    for 라벨, 글, 건수 in 서술항목:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        run(p, 라벨, size=9, bold=True)
        run(p, f"   전체 {건수}/{N}명 작성", size=8, color=GRAY)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Pt(14)
        p.paragraph_format.line_spacing = 1.4
        if 무응답(글):
            run(p, "무응답" if not 글 else f"'{글}'", size=9, color=GRAY, italic=True)
        else:
            run(p, 글, size=9.5)

    # ══ 6. 관련 이슈 ══
    관련 = [f for f in 집계["플래그"]
            if any(a["ID"] == r.ID for a in f["첨부"])]
    spacer(doc, 8)
    h(doc, "6. 이 응답이 근거로 들어간 항목", 14, ACCENT, before=10, after=5, rule=True)
    if not 관련:
        body_p(doc, "종합 리포트 3장의 어느 항목에도 이 응답이 근거로 인용되지 않았다.",
               size=9, color=GRAY)
    else:
        body_p(doc, f"종합 리포트 3장 {len(관련)}개 항목이 이 응답을 근거로 삼는다.",
               size=8.5, color=GRAY, after=4)
        datatable(doc, ["심각도", "항목", "근거 수치"],
                  [[f["심각도"], 자르기(f["제목"], 38), 자르기(f["근거"], 52)] for f in 관련],
                  [1300, 4200, 4137], sizes=(8.3, 8.3))

    doc.save(경로)
    return 경로


def 파일명(게임, r):
    안전 = 게임["게임명"].replace("/", "／").replace(":", "：")
    return f"「{안전}」_{r.ID}_응답카드.docx"
