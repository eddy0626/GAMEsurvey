# -*- coding: utf-8 -*-
"""게임별 종합 리포트 (12~15페이지)

참조/미스트월드_플레이테스트_결과리포트.docx 와 같은 장 구조:
  한눈에 보기 → 1.테스트 개요 → 2.정량 결과 → 3.핵심 발견
  → 4.세그먼트 분석 → 5.개선 권고 → 부록 A · B · C

3장 핵심 발견은 완전 자동 생성하지 않는다.
이슈 플래그에 걸린 항목만 카드로 나열하고 근거 수치와 응답자 인용을 붙인
초안까지만 만든다. 권고 문장 자리는 비워 둔다.
"""
import os

from docx.shared import Pt

from . import charts
from . import config as C
from .aggregate import 백분율, 찾기
from .kor import 이가, 자르기
from .docxkit import (ACCENT, GRAY, SEV, body_p, bullet, card, caption,
                      datatable, 문서만들기, finding, fixed_table, h, pic, run,
                      shade, spacer, white_borders, cell_margins)

_심각도순 = {"Critical": 0, "High": 1, "Medium": 2, "검토필요": 3, "강점": 4, "Low": 5}


def _KPI(doc, 지표):
    N = 지표["N"]
    최고 = 지표["최고축"]; 최저 = 지표["최저축"]
    가격최빈 = max(지표["적정가격"], key=lambda r: r["인원"], default=None)
    항목 = [
        ("6축 평균", f'{지표["육각총평균"]:.2f}' if 지표["육각총평균"] is not None else "-", "5점 만점"),
        ("최고 항목", 최고["이름"] if 최고 else "-", f'{최고["평균"]:.2f}' if 최고 else "-"),
        ("최저 항목", 최저["이름"] if 최저 else "-", f'{최저["평균"]:.2f}' if 최저 else "-"),
        ("NPS", str(지표["NPS"]["값"]) if 지표["NPS"]["값"] is not None else "-",
         f'추천 {지표["NPS"]["추천"]} · 중립 {지표["NPS"]["중립"]} · 비추천 {지표["NPS"]["비추천"]}'),
        ("구매 의향", f'{지표["구매의향률"]["비율"]}%', "정가 또는 구매 의향"),
        ("적정가 최빈", 가격최빈["답"] if 가격최빈 and 가격최빈["인원"] else "-",
         f'{가격최빈["인원"]}명' if 가격최빈 else "-"),
    ]
    t = fixed_table(doc, [1606] * 6)
    for i, (라벨, 값, 부) in enumerate(항목):
        c = t.rows[0].cells[i]
        white_borders(c); shade(c, "F1F6FB"); cell_margins(c, 120, 90, 120, 90)
        c.paragraphs[0].text = ""
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(c.paragraphs[0], 라벨, size=8, color=GRAY, bold=True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        run(p2, 값, size=11 if len(값) > 6 else 13.5, bold=True, color=ACCENT)
        p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        run(p3, 부, size=7.5, color=GRAY)
    spacer(doc, 10)


def 만들기(집계, 경로, 차트폴더):
    게임, 지표, N = 집계["게임"], 집계["지표"], 집계["지표"]["N"]
    doc = 문서만들기()

    # ── 표지 ──
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, C.사업명, size=9, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, f"{게임['게임명']} 플레이테스트 결과 리포트", size=25, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    run(p, f"개발사 {게임['개발사']}  ·  응답 {N}명  ·  작성 {C.기관명}", size=10.5, color=GRAY)

    if N == 0:
        h(doc, "응답이 아직 없습니다", 15, ACCENT)
        body_p(doc, "이 게임의 응답 시트에 응답이 한 건도 없어 리포트를 만들 수 없습니다. "
                    "테스트 마감 후 다시 생성해 주세요.", size=10)
        doc.save(경로)
        return 경로

    # ── 한눈에 보기 ──
    h(doc, "한눈에 보기", 15, ACCENT, before=4, after=6)
    _KPI(doc, 지표)

    body = card(doc, ACCENT, "F1F6FB")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    run(p, "결론 세 줄  (담당자 확정 필요)", size=10.5, bold=True)
    최고, 최저 = 지표["최고축"], 지표["최저축"]
    초안 = []
    if 최고:
        초안.append(f'가장 높은 항목은 {최고["이름"]} {최고["평균"]:.2f}점입니다. '
                    f'4~5점을 준 응답자가 {지표["육각"][[a["이름"] for a in 지표["육각"]].index(최고["이름"])]["긍정"]}명입니다.')
    if 최저:
        초안.append(f'가장 낮은 항목은 {최저["이름"]} {최저["평균"]:.2f}점입니다. '
                    f'진행도가 튜토리얼 단계에 머문 응답자가 {지표["진행도정체"]["인원"]}명'
                    f'({지표["진행도정체"]["비율"]}%)입니다.')
    초안.append(f'NPS {지표["NPS"]["값"]} · 유료 구매를 고려한 응답자 {지표["유료고려"]["비율"]}%. '
                f'개발 중 빌드를 짧게 본 결과라는 점을 감안해 읽어야 합니다.')
    for txt in 초안:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.35
        run(p, "· ", size=9.5, bold=True, color=ACCENT)
        run(p, txt, size=9.5)
    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    run(p, "위 세 줄은 수치에서 뽑은 초안입니다. 담당자가 문장을 확정해 주세요.",
        size=8.5, italic=True, color=GRAY)
    spacer(doc, 8)

    # ── 먼저 손대야 할 것 ──
    상위 = [f for f in 집계["플래그"] if f["심각도"] in ("Critical", "High")][:6]
    if 상위:
        h(doc, "먼저 손대야 할 것", 13, before=8, after=5, rule=True)
        datatable(doc, ["순위", "항목", "근거 수치", "심각도"],
                  [[i + 1, 자르기(f["제목"], 44), 자르기(f["근거"], 72), f["심각도"]]
                   for i, f in enumerate(상위)],
                  [700, 3300, 4300, 1337])
        spacer(doc, 4)
        body_p(doc, "심각도는 게임 유저 리서치에서 통용되는 4단계(Critical / High / Medium / Low)를 "
                    "적용했습니다. Critical은 진행이 막히거나 다수가 이탈한 문제, "
                    "High는 다음 빌드 안에 고치는 것이 좋은 문제입니다.",
               size=8.5, color=GRAY, after=4)

    # ── 1. 테스트 개요 ──
    doc.add_page_break()
    h(doc, "1. 테스트 개요", 16, ACCENT, before=0, after=6, rule=True)
    유형표 = {}
    for r in 집계["응답"]:
        유형표[r.유형] = 유형표.get(r.유형, 0) + 1
    datatable(doc, ["항목", "내용"], [
        ["대상 빌드", f"{게임['게임명']} 플레이테스트 빌드 (개발 중)"],
        ["방식", "현장 플레이 후 Google Forms 설문 (자기 기입식)"],
        ["표본", f"{N}명 — " + " · ".join(f'{v["답"]} {v["인원"]}명' for v in 지표["연령대"])],
        ["응답자 유형", " · ".join(f"{k} {v}명" for k, v in 유형표.items()) or "-"],
        ["성별", " · ".join(f'{v["답"]} {v["인원"]}명' for v in 지표["성별"]) or "-"],
        ["장르 선호", " · ".join(f'{v["답"]} {v["인원"]}명' for v in 지표["선호장르"][:4]) + " 외"],
        ["설문 문항", f"{len(집계['응답'][0].고유)}개 회사 문항 + 공통 22문항"],
        ["응답률", f'개선점 서술 {지표["응답률"]["개선점"]}/{N} · '
                   f'개발팀 메시지 {지표["응답률"]["개발팀메시지"]}/{N}'],
    ], [1900, 7737])
    spacer(doc, 8)

    body = card(doc, "E8A33D", "FFF4E5")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    run(p, "해석할 때 감안할 점", size=10, bold=True)
    한계 = [f"표본이 {N}명입니다. 비율(%)은 경향을 보는 용도이고, "
            f"1명 차이가 {round(100/N)}%p 안팎을 움직입니다. "
            "숫자보다 여러 사람이 같은 말을 한 항목에 무게를 두시는 편이 안전합니다."]
    완주 = 찾기(지표["진행도"], "준비된 분량을 모두 경험했다")["인원"]
    if 완주 == 0:
        한계.append("준비된 분량을 모두 경험한 응답자는 0명입니다. "
                    "중후반 콘텐츠 · 밸런스 · 장기 몰입도는 이번 데이터로 판단할 수 없습니다.")
    한계.append("교육 수강생 중심 표본이라 장르 선호가 고르게 섞였습니다. "
                "비선호층의 낮은 점수를 곧바로 게임 결함으로 읽으면 안 됩니다. 아래 4장에서 나눠 봤습니다.")
    for txt in 한계:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.3
        run(p, "· ", size=9, bold=True, color="B45309"); run(p, txt, size=9)
    spacer(doc, 6)

    # ── 2. 정량 결과 ──
    doc.add_page_break()
    h(doc, "2. 정량 결과", 16, ACCENT, before=0, after=6, rule=True)
    h(doc, "2-1. 게임 전체 평가 (육각형 6축)", 12, before=6, after=4)

    코드 = 게임["코드"]
    레이더 = os.path.join(차트폴더, f"{코드}_radar.png")
    막대 = os.path.join(차트폴더, f"{코드}_bars.png")
    퍼널 = os.path.join(차트폴더, f"{코드}_funnel.png")
    charts.레이더_게임([a["평균"] for a in 지표["육각"]], 레이더, 게임["게임명"], N)
    charts.막대_항목평균([a["이름"] for a in 지표["육각"] + 지표["추가"]],
                        [a["평균"] for a in 지표["육각"] + 지표["추가"]], 막대, N)
    charts.퍼널_진행도구매(지표["진행도"], 지표["구매의향"], 퍼널, N)

    pic(doc, 레이더, 11.4)
    caption(doc, f'그림 1. 육각형 평가 — 6축 평균 {지표["육각총평균"]:.2f} / 5점 만점')
    pic(doc, 막대, 16.4)
    caption(doc, "그림 2. 8개 항목 평균 (육각형 6축 + 진행 흐름 · 기술 안정성)")

    datatable(doc, ["항목", "평균", "표준편차", "긍정 (4~5점)", "부정 (1~2점)"],
              [[a["이름"],
                f'{a["평균"]:.2f}' if a["평균"] is not None else "-",
                f'{a["표준편차"]:.2f}' if a["표준편차"] is not None else "-",
                f'{a["긍정"]}명', f'{a["부정"]}명'] for a in 지표["육각"] + 지표["추가"]],
              [3037, 1500, 1600, 1750, 1750])
    spacer(doc, 6)
    편차있는 = [a for a in 지표["육각"] + 지표["추가"] if a["표준편차"] is not None]
    if 편차있는:
        큰 = max(편차있는, key=lambda a: a["표준편차"])
        작은 = min(편차있는, key=lambda a: a["표준편차"])
        body_p(doc, f'표준편차가 큰 항목은 사람마다 평가가 갈렸다는 뜻입니다. '
                    f'{이가(큰["이름"])} {큰["표준편차"]:.2f}로 가장 크게 갈렸고, '
                    f'{이가(작은["이름"])} {작은["표준편차"]:.2f}로 가장 의견이 모였습니다.',
               size=9, color=GRAY)

    h(doc, "2-2. 진행도와 구매 의향", 12, before=12, after=4)
    pic(doc, 퍼널, 16.4)
    caption(doc, "그림 3. 진행도(좌) · 스팀 구매 의향(우)")
    bullet(doc, f'준비된 분량을 모두 경험한 응답자는 {완주}명입니다.')
    bullet(doc, f'튜토리얼 단계에 머문 응답자가 {지표["진행도정체"]["인원"]}명'
                f'({지표["진행도정체"]["비율"]}%)입니다.')
    bullet(doc, f'유료 구매를 고려한 응답자는 {지표["유료고려"]["인원"]}명'
                f'({지표["유료고려"]["비율"]}%)입니다. '
                f'정가 구매 {찾기(지표["구매의향"], C.구매선택지["정가"])["인원"]}명 · '
                f'구매 의향 {찾기(지표["구매의향"], C.구매선택지["의향"])["인원"]}명 · '
                f'할인 시 {찾기(지표["구매의향"], C.구매선택지["할인"])["인원"]}명.')

    h(doc, "2-3. 추천 의향 (NPS)", 12, before=12, after=4)
    nps = 지표["NPS"]
    datatable(doc, ["구분", "점수대", "인원", "비율"],
              [["추천 (Promoter)", "9~10점", f'{nps["추천"]}명', f'{nps["추천율"]}%'],
               ["중립 (Passive)", "7~8점", f'{nps["중립"]}명', f'{nps["중립율"]}%'],
               ["비추천 (Detractor)", "0~6점", f'{nps["비추천"]}명', f'{nps["비추천율"]}%'],
               ["NPS = 추천% − 비추천%", "", f'{nps["값"]}', ""]],
              [3200, 2000, 2200, 2237])
    spacer(doc, 5)
    body_p(doc, f'NPS는 추천 비율에서 비추천 비율을 뺀 값입니다. '
                f'개발 중 빌드를 짧게 본 결과라 절대값보다 분포를 보는 편이 맞습니다. '
                f'평균 추천 점수는 {nps["평균"]}점입니다.', size=9.5)

    h(doc, "2-4. 적정 가격", 12, before=12, after=4)
    최빈인원 = max((r["인원"] for r in 지표["적정가격"]), default=0)
    datatable(doc, ["가격대", "인원", "비율", "비고"],
              [[k, f'{찾기(지표["적정가격"], k)["인원"]}명', f'{찾기(지표["적정가격"], k)["비율"]}%',
                "최빈" if 찾기(지표["적정가격"], k)["인원"] == 최빈인원 and 최빈인원 else ""]
               for k in C.가격구간],
              [3000, 1600, 1600, 3437])
    spacer(doc, 5)
    만원이상 = sum(찾기(지표["적정가격"], k)["인원"] for k in C.가격구간[2:])
    body_p(doc, f'1만 원 이상을 적정가로 본 응답자가 {만원이상}명({백분율(만원이상, N)}%)입니다.',
           size=9.5)

    # ── 3. 핵심 발견 ──
    doc.add_page_break()
    h(doc, "3. 핵심 발견", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "이슈 플래그 규칙에 걸린 항목을 심각도 순으로 나열했습니다. "
                "각 카드에 정량 근거와 응답자 원문을 함께 실었습니다. "
                "응답자는 P01 형식으로 익명 처리했습니다.", size=9, color=GRAY, after=4)
    body = card(doc, "E8A33D", "FFF4E5")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    run(p, "담당자가 할 일  ", size=9, bold=True, color="B45309")
    run(p, "아래는 자동으로 뽑은 '이슈 후보'입니다. 어떤 것을 리포트에 올릴지 고르고, "
           "각 카드의 권고 자리를 채운 뒤 나머지는 지워 주세요. "
           "규칙은 무엇이 걸렸는지까지만 알려 줄 뿐 권고 문장을 대신 쓰지 않습니다.", size=9)
    spacer(doc, 8)

    플래그 = sorted(집계["플래그"], key=lambda f: _심각도순.get(f["심각도"], 9))
    for i, f in enumerate(플래그):
        인용 = []
        본 = set()
        for a in f["첨부"]:
            키 = (a["ID"], a["글"][:24])
            if 키 in 본:
                continue
            본.add(키)
            인용.append((a["ID"], a["글"], a.get("라벨", "")))
            if len(인용) >= 4:
                break
        finding(doc, f["심각도"], f'F{i+1}', f["제목"], f["근거"], 인용, 권고자리=True)

    if not 플래그:
        body_p(doc, "임계값에 걸린 항목이 없습니다. Config 의 임계값을 조정하거나 "
                    "응답이 더 쌓인 뒤 다시 생성해 주세요.", size=9.5, color=GRAY)

    # ── 4. 세그먼트 ──
    doc.add_page_break()
    h(doc, "4. 누가 어떻게 갈렸는가", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "평균값만 보면 놓치는 부분이라 응답자를 나눠 봤습니다.", size=10, after=10)

    세그 = 집계["세그먼트"]
    h(doc, "4-1. 장르 선호에 따른 분화", 12, before=4, after=4)
    datatable(doc, ["집단", "인원", "6축 평균", "NPS 평균", "진행 흐름", "튜토리얼 단계 정체"],
              [[g["라벨"], f'{g["인원"]}명',
                f'{g["육각평균"]:.2f}' if g["육각평균"] is not None else "-",
                f'{g["NPS평균"]:.1f}' if g["NPS평균"] is not None else "-",
                f'{g["진행흐름"]:.2f}' if g["진행흐름"] is not None else "-",
                f'{g["정체"]}명'] for g in 세그["장르"]],
              [2600, 1000, 1500, 1400, 1400, 1737])
    spacer(doc, 6)
    body_p(doc, f'코어 장르는 {" · ".join(세그["코어장르"])}로 잡았습니다. '
                f'(Config 의 게임_코어장르에서 조정할 수 있습니다)', size=8.5, color=GRAY)

    if 세그["튜토리얼"]:
        h(doc, "4-2. 튜토리얼 이해도에 따른 분화", 12, before=10, after=4)
        datatable(doc, ["집단", "인원", "6축 평균", "NPS 평균", "진행 흐름", "출시 기대감"],
                  [[g["라벨"], f'{g["인원"]}명',
                    f'{g["육각평균"]:.2f}' if g["육각평균"] is not None else "-",
                    f'{g["NPS평균"]:.1f}' if g["NPS평균"] is not None else "-",
                    f'{g["진행흐름"]:.2f}' if g["진행흐름"] is not None else "-",
                    f'{g["출시기대감"]:.2f}' if g["출시기대감"] is not None else "-"]
                   for g in 세그["튜토리얼"]["행"]],
                  [2600, 1000, 1500, 1400, 1500, 1637])
    else:
        h(doc, "4-2. 본편 진입 여부에 따른 분화", 12, before=10, after=4)
        body_p(doc, "이 게임은 튜토리얼 이해도를 객관식으로 묻지 않아 진행도로 나눴습니다.",
               size=8.5, color=GRAY, after=4)
        datatable(doc, ["집단", "인원", "6축 평균", "NPS 평균", "진행 흐름", "출시 기대감"],
                  [[g["라벨"], f'{g["인원"]}명',
                    f'{g["육각평균"]:.2f}' if g["육각평균"] is not None else "-",
                    f'{g["NPS평균"]:.1f}' if g["NPS평균"] is not None else "-",
                    f'{g["진행흐름"]:.2f}' if g["진행흐름"] is not None else "-",
                    f'{g["출시기대감"]:.2f}' if g["출시기대감"] is not None else "-"]
                   for g in 세그["진행"]],
                  [2600, 1000, 1500, 1400, 1500, 1637])
    spacer(doc, 8)

    # ── 5. 개선 권고 ──
    doc.add_page_break()
    h(doc, "5. 개선 권고 정리", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "3장에서 채택한 항목을 시기별로 나눠 적는 표입니다. "
                "아래는 심각도로 자동 배치한 초안이니, 3장을 확정한 뒤 손봐 주세요.",
           size=9, color=GRAY, after=6)
    시기 = {"Critical": "다음 빌드", "High": "단기", "Medium": "중기",
            "검토필요": "중기", "강점": "유지", "Low": "출시 전"}
    행5 = [[시기.get(f["심각도"], "중기"), f'F{i+1}', 자르기(f["제목"], 42), "(담당자 작성)"]
           for i, f in enumerate(플래그) if f["심각도"] != "강점"]
    if 행5:
        datatable(doc, ["시기", "항목", "무엇을", "구체 작업 · 기대 효과"], 행5,
                  [1100, 900, 4200, 3437], sizes=(8.5, 8.3))
    강점들 = [f for f in 플래그 if f["심각도"] == "강점"]
    if 강점들:
        spacer(doc, 8)
        body = card(doc, "2E9E5B", "EAF7EF")
        p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
        run(p, "지켜야 할 것", size=10, bold=True)
        for f in 강점들:
            p = body.add_paragraph()
            p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.3
            run(p, "· ", size=9, bold=True, color="2E9E5B")
            run(p, f'{f["제목"]} — {f["근거"]}', size=9)

    # ── 부록 A ──
    doc.add_page_break()
    h(doc, "부록 A. 문항별 응답 집계", 16, ACCENT, before=0, after=6, rule=True)
    for 항 in 집계["문항집계"]:
        if 항["유형"] == "서술형":
            continue
        제목 = 항["문항"]
        if 항["유형"] == "선형배율" and 항["통계"]:
            제목 += f'   (평균 {항["통계"]["평균"]:.2f} · 표준편차 ' \
                    f'{항["통계"]["표준편차"]:.2f})' if 항["통계"]["표준편차"] is not None \
                    else f'   (평균 {항["통계"]["평균"]:.2f})'
        h(doc, 제목, 10.5, before=9, after=3)
        datatable(doc, ["응답", "인원", "비율"],
                  [[r["답"], f'{r["인원"]}명', f'{r["비율"]}%'] for r in 항["집계"]],
                  [6437, 1600, 1600], sizes=(8.5, 8.5))

    # ── 부록 B ──
    doc.add_page_break()
    h(doc, "부록 B. 자유 서술 원문", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "응답자가 직접 작성한 내용을 손대지 않고 그대로 실었습니다. "
                "무응답과 '없음'은 제외했습니다.", size=9, color=GRAY, after=8)
    for 묶음 in 집계["서술모음"]:
        if not 묶음["건수"]:
            continue
        h(doc, f'{묶음["라벨"]}  ({묶음["건수"]}건)', 11, before=10, after=4)
        for 항목 in 묶음["항목"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Pt(6)
            p.paragraph_format.line_spacing = 1.3
            run(p, f'{항목["ID"]}  ', size=8.5, bold=True, color=ACCENT)
            run(p, 항목["글"], size=9)

    # ── 부록 C ──
    doc.add_page_break()
    h(doc, "부록 C. 응답자별 평가 점수", 16, ACCENT, before=0, after=6, rule=True)
    행C = []
    for r in 집계["응답"]:
        행C.append([r.ID, r.연령대, (r.선호장르[0] if r.선호장르 else "-")]
                   + [v if v is not None else "-" for v in r.육각]
                   + [r.진행흐름 if r.진행흐름 is not None else "-",
                      r.기술안정성 if r.기술안정성 is not None else "-",
                      r.추천의향 if r.추천의향 is not None else "-"])
    datatable(doc, ["ID", "연령", "주 장르", "재미", "조작", "아트", "몰입", "완성", "기대",
                    "흐름", "안정", "NPS"], 행C,
              [620, 700, 1800, 620, 620, 620, 620, 620, 620, 620, 620, 620], sizes=(8, 8))
    spacer(doc, 6)
    body_p(doc, "응답자 실명은 이 문서 어디에도 없습니다. "
                "이름과 ID 의 대응표는 마스터 시트의 접근 제한 시트에만 있습니다.",
           size=8.5, color=GRAY)

    doc.save(경로)
    return 경로


def 파일명(게임):
    안전 = 게임["게임명"].replace("/", "／").replace(":", "：")
    return f"{안전}_플레이테스트_결과리포트.docx"
