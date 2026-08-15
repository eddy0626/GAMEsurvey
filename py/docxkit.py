# -*- coding: utf-8 -*-
"""docx 저수준 헬퍼

참조/gen_report.py 에서 그대로 가져왔다. 그 파일이 만든 미스트월드 리포트가
목표 산출물이므로, 서식을 새로 만들지 않고 검증된 것을 재사용한다.
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

FONT = "맑은 고딕"
ACCENT = "2E74B5"
GRAY = "667085"
SEV = {
    "Critical": ("C0392B", "FDECEA"),
    "High":     ("E8760D", "FEF3E6"),
    "Medium":   ("C9A227", "FDF8E3"),
    "Low":      ("5B7C99", "EEF3F7"),
    "검토필요":  ("5B7C99", "EEF3F7"),
    "강점":      ("2E9E5B", "EAF7EF"),
}


def 문서만들기(가로=21.0, 세로=29.7, 여백=2.0):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(가로), Cm(세로)
    sec.left_margin = sec.right_margin = Cm(여백)
    sec.top_margin = sec.bottom_margin = Cm(여백)
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for p0 in list(doc.paragraphs):
        p0._element.getparent().remove(p0._element)
    return doc


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def white_borders(cell, color="FFFFFF"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)
        b.append(el)
    tcPr.append(b)


def cell_margins(cell, t, l, b_, r):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for name, val in (("top", t), ("left", l), ("bottom", b_), ("right", r)):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(dxa)); w.set(qn("w:type"), "dxa")
    tcPr.insert(0, w)


def run(par, text, size=10, bold=False, color=None, italic=False):
    r = par.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ("w:ascii", "w:eastAsia", "w:hAnsi"):
        rf.set(qn(a), FONT)
    return r


def spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt); p.paragraph_format.space_before = Pt(0)
    return p


def fixed_table(doc, widths, total=None):
    total = total or sum(widths)
    t = doc.add_table(rows=1, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    tblPr = t._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    w = OxmlElement("w:tblW"); w.set(qn("w:w"), str(total)); w.set(qn("w:type"), "dxa")
    tblPr.append(w)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed"); tblPr.append(lay)
    grid = t._tbl.find(qn("w:tblGrid"))
    for gc in list(grid):
        grid.remove(gc)
    for x in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(x)); grid.append(gc)
    for c, x in zip(t.rows[0].cells, widths):
        c.width = Emu(int(x * 635)); set_cell_width(c, x)
    return t


def card(doc, bar_color, fill_color, width=9637, bar=170):
    t = fixed_table(doc, [bar, width - bar], total=width)
    b, body = t.rows[0].cells
    trPr = t.rows[0]._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))
    for c in (b, body):
        white_borders(c)
    shade(b, bar_color); shade(body, fill_color)
    cell_margins(b, 0, 0, 0, 0); cell_margins(body, 130, 200, 130, 200)
    body.paragraphs[0].text = ""
    bp = b.paragraphs[0]; bp.paragraph_format.space_after = Pt(0)
    return body


def h(doc, text, size, color=None, before=14, after=4, rule=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    run(p, text, size=size, bold=True, color=color)
    if rule:
        pPr = p._p.get_or_add_pPr()
        pbd = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
        bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6")
        bt.set(qn("w:space"), "4"); bt.set(qn("w:color"), "D6DEE7")
        pbd.append(bt); pPr.append(pbd)
    return p


def body_p(doc, text, size=10, color=None, after=6, indent=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    run(p, text, size=size, color=color, bold=bold)
    return p


def bullet(doc, text, size=9.5, indent=10, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.line_spacing = 1.3
    run(p, "· ", size=size, color=color, bold=True)
    run(p, text, size=size)
    return p


def datatable(doc, header, rows, widths, sizes=(8.8, 8.8), highlight=None):
    t = fixed_table(doc, widths)
    t.style = "Table Grid"
    for i, txt in enumerate(header):
        c = t.rows[0].cells[i]
        c.paragraphs[0].text = ""
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(c.paragraphs[0], str(txt), size=sizes[0], bold=True)
        shade(c, "DCE9F5")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].paragraphs[0].text = ""
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.2
            run(cells[i].paragraphs[0], str(v), size=sizes[1],
                bold=(highlight is not None and i == highlight))
    return t


def pic(doc, 경로, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(경로, width=Cm(width_cm))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run(p, text, size=8.5, color=GRAY, italic=True)


def finding(doc, sev, code, title, evidence, quotes, action=None, 권고자리=False):
    """3장 발견 카드.
    권고자리=True 면 권고 문장을 쓰지 않고 담당자가 채울 자리만 남긴다.
    (규칙만으로 권고를 생성하지 않는다는 원칙)"""
    fg, bg = SEV.get(sev, SEV["Low"])
    body = card(doc, fg, bg)
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run(p, f"[{sev}]", size=8.5, bold=True, color=fg)
    run(p, f"  {code}. {title}", size=10.5, bold=True)

    p = body.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run(p, "근거  ", size=8.5, bold=True, color=GRAY)
    run(p, evidence, size=9)

    for who, quote, label in quotes:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run(p, f"“{quote}”", size=8.8, italic=True, color="3A4A5A")
        run(p, f"  — {who}" + (f" · {label}" if label else ""), size=8, color=GRAY)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    run(p, "권고  ", size=8.5, bold=True, color=fg)
    if 권고자리:
        run(p, "여기에 권고를 작성하세요. "
               "위 근거와 인용을 보고 채택 여부를 정한 뒤, 무엇을 어떻게 고칠지 적습니다.",
            size=9, italic=True, color=GRAY)
    else:
        run(p, action or "", size=9)
    spacer(doc, 6)
    return body
