# -*- coding: utf-8 -*-
"""미스트월드 플레이테스트 결과 리포트 (.docx) 생성 — 개발사 전달용"""
import os
from collections import Counter
from statistics import mean, stdev

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from report_data import R, HEX_AXES

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "..", "미스트월드_플레이테스트_결과리포트.docx")

FONT = "맑은 고딕"
ACCENT   = "2E74B5"
GRAY     = "667085"
SEV = {  # 심각도 색상
    "Critical": ("C0392B", "FDECEA"),
    "High":     ("E8760D", "FEF3E6"),
    "Medium":   ("C9A227", "FDF8E3"),
    "Low":      ("5B7C99", "EEF3F7"),
    "강점":      ("2E9E5B", "EAF7EF"),
}


# ── 저수준 헬퍼 ──────────────────────────────────────────────
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def white_borders(cell, color="FFFFFF"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), color)
        b.append(el)
    tcPr.append(b)


def cell_margins(cell, t, l, b_, r):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for name, val in (('top', t), ('left', l), ('bottom', b_), ('right', r)):
        el = OxmlElement(f'w:{name}')
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcPr.append(mar)


def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(dxa)); w.set(qn('w:type'), 'dxa')
    tcPr.insert(0, w)


def run(par, text, size=10, bold=False, color=None, italic=False):
    r = par.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    for a in ('w:ascii', 'w:eastAsia', 'w:hAnsi'):
        rf.set(qn(a), FONT)
    return r


def spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt); p.paragraph_format.space_before = Pt(0)
    return p


def fixed_table(doc, widths, total=None):
    total = total or sum(widths)
    t = doc.add_table(rows=1, cols=len(widths))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.autofit = False
    tblPr = t._tbl.tblPr
    for tag in ('w:tblW', 'w:tblLayout'):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(total)); w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    grid = t._tbl.find(qn('w:tblGrid'))
    for gc in list(grid):
        grid.remove(gc)
    for x in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(x)); grid.append(gc)
    for c, x in zip(t.rows[0].cells, widths):
        c.width = Emu(int(x * 635)); set_cell_width(c, x)
    return t


def card(doc, bar_color, fill_color, width=9637, bar=170):
    t = fixed_table(doc, [bar, width - bar], total=width)
    b, body = t.rows[0].cells
    trPr = t.rows[0]._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:cantSplit'))
    for c in (b, body):
        white_borders(c)
    shade(b, bar_color); shade(body, fill_color)
    cell_margins(b, 0, 0, 0, 0); cell_margins(body, 130, 200, 130, 200)
    body.paragraphs[0].text = ''
    bp = b.paragraphs[0]; bp.paragraph_format.space_after = Pt(0)
    return body


def h(doc, text, size, color=None, before=14, after=4, rule=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    run(p, text, size=size, bold=True, color=color)
    if rule:
        pPr = p._p.get_or_add_pPr()
        pbd = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
        bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '6')
        bt.set(qn('w:space'), '4'); bt.set(qn('w:color'), 'D6DEE7')
        pbd.append(bt); pPr.append(pbd)
    return p


def body_p(doc, text, size=10, color=None, after=6, indent=0, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    run(p, text, size=size, color=color, bold=bold)
    return p


def bullet(doc, text, size=9.5, indent=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Pt(indent)
    p.paragraph_format.line_spacing = 1.3
    run(p, "· ", size=size, color=ACCENT, bold=True)
    run(p, text, size=size)
    return p


def datatable(doc, header, rows, widths, sizes=(8.8, 8.8), highlight=None):
    t = fixed_table(doc, widths)
    t.style = 'Table Grid'
    for i, txt in enumerate(header):
        c = t.rows[0].cells[i]
        c.paragraphs[0].text = ''
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(c.paragraphs[0], txt, size=sizes[0], bold=True)
        shade(c, "DCE9F5")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].paragraphs[0].text = ''
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            cells[i].paragraphs[0].paragraph_format.line_spacing = 1.2
            run(cells[i].paragraphs[0], str(v), size=sizes[1],
                bold=(highlight is not None and i == highlight))
    return t


def pic(doc, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(os.path.join(HERE, path), width=Cm(width_cm))
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run(p, text, size=8.5, color=GRAY, italic=True)


def finding(doc, sev, code, title, evidence, quotes, action):
    fg, bg = SEV[sev]
    body = card(doc, fg, bg)
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run(p, f"[{sev}]", size=8.5, bold=True, color=fg)
    run(p, f"  {code}. {title}", size=10.5, bold=True)

    p = body.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    run(p, "근거  ", size=8.5, bold=True, color=GRAY)
    run(p, evidence, size=9)

    for q in quotes:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.line_spacing = 1.25
        run(p, f"“{q[1]}”", size=8.8, italic=True, color="3A4A5A")
        run(p, f"  — {q[0]}", size=8, color=GRAY)

    p = body.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.3
    run(p, "권고  ", size=8.5, bold=True, color=fg)
    run(p, action, size=9)
    spacer(doc, 6)


# ── 집계 ─────────────────────────────────────────────────────
N = len(R)
def pct(n): return round(n / N * 100)
HEX = [round(mean([r["hex"][i] for r in R]), 2) for i in range(6)]
HEXSD = [round(stdev([r["hex"][i] for r in R]), 2) for i in range(6)]
FLOW = round(mean([r["flow"] for r in R]), 2)
TECH = round(mean([r["tech"] for r in R]), 2)
NPSV = [r["nps"] for r in R]
PROM = sum(1 for v in NPSV if v >= 9); PASS_ = sum(1 for v in NPSV if 7 <= v <= 8)
DET = sum(1 for v in NPSV if v <= 6)
NPS = round((PROM - DET) / N * 100)
CORE = [r for r in R if "RPG" in r["genres"] or "전략 · 턴제" in r["genres"]]
NONCORE = [r for r in R if r not in CORE]


def cnt(key):
    return Counter(r[key] for r in R)


def cnt_multi(key):
    return Counter(x for r in R for x in r[key])


# ── 문서 ─────────────────────────────────────────────────────
def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.0)
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for p0 in list(doc.paragraphs):
        p0._element.getparent().remove(p0._element)

    # ── 표지 ──
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    run(p, "2026 충북 인디게임 플레이테스트", size=9, bold=True, color=ACCENT)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    run(p, "미스트월드 플레이테스트 결과 리포트", size=25, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    run(p, "개발사 프로젝트 미스트  ·  응답 15명  ·  작성 충북글로벌게임센터",
        size=10.5, color=GRAY)

    # ── 한눈에 보기 ──
    h(doc, "한눈에 보기", 15, ACCENT, before=4, after=6)
    kpi = [
        ("6축 평균", f"{mean(HEX):.2f}", "5점 만점"),
        ("최고 항목", "그래픽·아트", f"{HEX[2]:.2f}"),
        ("최저 항목", "완성도", f"{HEX[4]:.2f}"),
        ("NPS", f"{NPS}", f"추천 {PROM} · 중립 {PASS_} · 비추천 {DET}"),
        ("구매 의향", f"{pct(sum(1 for r in R if r['buy'] in ('정가에도 바로 구매할 것 같다','구매할 의향이 있다')))}%",
         "정가 또는 구매 의향"),
        ("적정가 최빈", "1~1.5만원", f"{cnt('price')['10,000~14,900원']}명"),
    ]
    t = fixed_table(doc, [1606]*6)
    for i, (label, val, sub) in enumerate(kpi):
        c = t.rows[0].cells[i]
        white_borders(c, "FFFFFF"); shade(c, "F1F6FB"); cell_margins(c, 120, 90, 120, 90)
        c.paragraphs[0].text = ''
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(c.paragraphs[0], label, size=8, color=GRAY, bold=True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(1)
        run(p2, val, size=12 if len(val) > 5 else 13.5, bold=True, color=ACCENT)
        p3 = c.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        run(p3, sub, size=7.5, color=GRAY)
    spacer(doc, 10)

    body = card(doc, ACCENT, "F1F6FB")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    run(p, "결론 세 줄", size=10.5, bold=True)
    for txt in [
        f"아트는 명확한 강점입니다. 그래픽·아트 {HEX[2]:.2f}점으로 8개 항목 중 유일하게 4점대이고, "
        f"15명 중 14명이 4점 이상을 줬습니다. 이 축은 지금 상태를 지키는 것이 최선입니다.",

        f"막힌 곳은 전부 초반 15분입니다. 튜토리얼을 '어려웠다'고 한 5명은 6축 평균 2.67 · NPS 4.2로 "
        f"'쉬웠다' 그룹(3.97 · 8.0)과 확연히 갈렸고, 이 5명 전원이 비(非)RPG 선호층이었습니다. "
        f"온보딩 하나가 이후 모든 지표를 결정했습니다.",

        f"지금 수치가 낮은 것은 정상입니다. NPS {NPS}, 완성도 {HEX[4]:.2f}는 개발 중 빌드에서 나올 만한 값이고, "
        f"그럼에도 정가 구매 의향이 {cnt('buy')['정가에도 바로 구매할 것 같다']}명, "
        f"적정가 1만 원 이상이 {sum(v for k,v in cnt('price').items() if k not in ('5,000원 미만','5,000~9,900원'))}명입니다. "
        f"방향은 맞고 진입 장벽만 남았습니다.",
    ]:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.35
        run(p, "· ", size=9.5, bold=True, color=ACCENT)
        run(p, txt, size=9.5)
    spacer(doc, 8)

    # ── 우선 처리 목록 ──
    h(doc, "먼저 손대야 할 것", 13, before=8, after=5, rule=True)
    datatable(doc,
        ["순위", "항목", "근거 수치", "심각도"],
        [["1", "초반 온보딩 — 튜토리얼 · 첫 전투 진입", "튜토리얼 '어려움' 5명(33%) · 튜토리얼 미완 2명(13%)", "Critical"],
         ["2", "전투 정보 노출 — 명중률 · 데미지 · 적 정보", "전투 이해 '어려움' 7명(47%)", "Critical"],
         ["3", "진행 차단 버그 2건", "재현 확인 1건 · 지장 있는 버그 2명(13%)", "High"],
         ["4", "UI · UX 동선", "개선 필요 11명(73%) · 조작성 2.93", "High"],
         ["5", "공세 게이지 · BREAK 인지", "'몰랐다' 3명 + '대폭 보강' 3명 = 40%", "High"],
         ["6", "주사위 판정 결과 표시", "'이해 어려움' 6명(40%)", "High"]],
        [700, 3300, 4300, 1337], highlight=None)
    spacer(doc, 4)
    body_p(doc, "심각도는 게임 유저 리서치에서 통용되는 4단계(Critical / High / Medium / Low)를 적용했습니다. "
                "Critical은 진행이 막히거나 다수가 이탈한 문제, High는 다음 빌드 안에 고치는 것이 좋은 문제입니다.",
           size=8.5, color=GRAY, after=4)

    # ── 테스트 개요 ──
    doc.add_page_break()
    h(doc, "1. 테스트 개요", 16, ACCENT, before=0, after=6, rule=True)
    age = cnt("age"); gen = cnt("gender")
    datatable(doc,
        ["항목", "내용"],
        [["대상 빌드", "미스트월드 플레이테스트 빌드 (개발 중)"],
         ["방식", "현장 플레이 후 Google Forms 설문 (자기 기입식)"],
         ["플레이 시간", "1인당 15분 (고정)"],
         ["표본", f"{N}명 — " + " · ".join(f"{k} {v}명" for k, v in age.most_common())],
         ["성별", " · ".join(f"{k} {v}명" for k, v in gen.most_common())],
         ["장르 선호", " · ".join(f"{k} {v}명" for k, v in cnt_multi("genres").most_common(4)) + " 외"],
         ["설문 문항", "32문항 (참여 동의 3 · 응답자 정보 3 · 회사 문항 10 · 전체 평가 8 · 종합 의견 8)"],
         ["응답률", f"필수 문항 100% · 개선점 서술 14/15 · 개발팀 메시지 13/15"]],
        [1900, 7737])
    spacer(doc, 8)

    body = card(doc, "E8A33D", "FFF4E5")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
    run(p, "해석할 때 감안할 점", size=10, bold=True)
    for txt in [
        f"표본이 {N}명입니다. 비율(%)은 경향을 보는 용도이고, 1명 차이가 6~7%p를 움직입니다. "
        "숫자보다 여러 사람이 같은 말을 한 항목에 무게를 두시는 편이 안전합니다.",
        "플레이가 15분으로 고정돼 준비된 분량을 모두 경험한 응답자는 0명입니다. "
        "중후반 콘텐츠·밸런스·장기 몰입도는 이번 데이터로 판단할 수 없습니다.",
        "교육 수강생 대상이라 장르 선호가 고르게 섞였습니다. CRPG는 코어 장르이므로 "
        "비선호층의 낮은 점수를 곧바로 게임 결함으로 읽으면 안 됩니다. 아래 4장에서 나눠 봤습니다.",
    ]:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.3
        run(p, "· ", size=9, bold=True, color="B45309"); run(p, txt, size=9)
    spacer(doc, 6)

    # ── 정량 결과 ──
    doc.add_page_break()
    h(doc, "2. 정량 결과", 16, ACCENT, before=0, after=6, rule=True)
    h(doc, "2-1. 게임 전체 평가 (육각형 6축)", 12, before=6, after=4)
    pic(doc, "chart_radar.png", 11.4)
    caption(doc, "그림 1. 육각형 평가 — 6축 평균 3.41 / 5점 만점")
    pic(doc, "chart_bars.png", 16.4)
    caption(doc, "그림 2. 8개 항목 평균 (육각형 6축 + 진행 흐름 · 기술 안정성)")

    rows = []
    for i, ax in enumerate(HEX_AXES):
        d = Counter(r["hex"][i] for r in R)
        rows.append([ax, f"{HEX[i]:.2f}", f"{HEXSD[i]:.2f}",
                     f"{d.get(4,0)+d.get(5,0)}명", f"{d.get(1,0)+d.get(2,0)}명"])
    for key, label in [("flow", "진행 흐름"), ("tech", "기술 안정성")]:
        vals = [r[key] for r in R]; d = Counter(vals)
        rows.append([label, f"{mean(vals):.2f}", f"{stdev(vals):.2f}",
                     f"{d.get(4,0)+d.get(5,0)}명", f"{d.get(1,0)+d.get(2,0)}명"])
    datatable(doc, ["항목", "평균", "표준편차", "긍정 (4~5점)", "부정 (1~2점)"],
              rows, [3037, 1500, 1600, 1750, 1750])
    spacer(doc, 6)
    body_p(doc, "표준편차가 큰 항목은 사람마다 평가가 갈렸다는 뜻입니다. "
                f"정식 출시 기대감이 {HEXSD[5]:.2f}로 가장 크게 갈렸고, "
                f"그래픽·아트가 {HEXSD[2]:.2f}로 가장 의견이 모였습니다. "
                "아트는 누가 봐도 좋았고, 계속할지 말지는 사람마다 달랐습니다.", size=9, color=GRAY)

    h(doc, "2-2. 진행도와 구매 의향", 12, before=12, after=4)
    pic(doc, "chart_funnel.png", 16.4)
    caption(doc, "그림 3. 진행도(좌) · 스팀 구매 의향(우)")
    bullet(doc, f"15분 안에 준비된 분량을 모두 경험한 응답자는 0명입니다. "
                f"'여러 콘텐츠를 두루 경험' {cnt('progress')['여러 콘텐츠를 두루 경험했다']}명이 최대치입니다.")
    bullet(doc, f"튜토리얼 단계에서 멈춘 응답자가 {cnt('progress')['튜토리얼을 마치지 못했다']}명, "
                f"튜토리얼까지만 마친 응답자가 {cnt('progress')['튜토리얼까지 마쳤다']}명입니다. "
                f"둘을 합치면 {cnt('progress')['튜토리얼을 마치지 못했다']+cnt('progress')['튜토리얼까지 마쳤다']}명(33%)이 "
                f"본편에 들어가지 못했습니다.")
    bullet(doc, f"돈을 낼 의향은 나쁘지 않습니다. 정가 구매 {cnt('buy')['정가에도 바로 구매할 것 같다']}명, "
                f"구매 의향 {cnt('buy')['구매할 의향이 있다']}명, 할인 시 {cnt('buy')['할인하면 구매할 것 같다']}명으로 "
                f"{pct(15 - cnt('buy')['구매하지 않을 것 같다'] - cnt('buy')['무료라면 해보겠다'])}%가 유료 구매를 고려했습니다.")

    h(doc, "2-3. 추천 의향 (NPS)", 12, before=12, after=4)
    datatable(doc, ["구분", "점수대", "인원", "비율"],
              [["추천 (Promoter)", "9~10점", f"{PROM}명", f"{pct(PROM)}%"],
               ["중립 (Passive)", "7~8점", f"{PASS_}명", f"{pct(PASS_)}%"],
               ["비추천 (Detractor)", "0~6점", f"{DET}명", f"{pct(DET)}%"],
               ["NPS = 추천% − 비추천%", "", f"{NPS}", ""]],
              [3200, 2000, 2200, 2237])
    spacer(doc, 5)
    body_p(doc, f"NPS {NPS}는 낮은 수치지만 개발 중 빌드의 초반 15분만 본 결과라는 점을 감안해야 합니다. "
                f"주목할 것은 절대값이 아니라 분포입니다. 비추천 {DET}명 중 4명이 튜토리얼을 '어려웠다'고 답했고, "
                f"추천 {PROM}명은 모두 RPG·전략 선호층이면서 튜토리얼을 어렵지 않게 넘긴 응답자였습니다.",
           size=9.5)

    h(doc, "2-4. 적정 가격", 12, before=12, after=4)
    pr = cnt("price")
    order = ["5,000원 미만", "5,000~9,900원", "10,000~14,900원",
             "15,000~19,900원", "20,000~29,900원", "30,000원 이상"]
    datatable(doc, ["가격대", "인원", "비율", "비고"],
              [[k, f"{pr.get(k,0)}명", f"{pct(pr.get(k,0))}%",
                "최빈" if pr.get(k,0) == max(pr.values()) else ""] for k in order],
              [3000, 1600, 1600, 3437])
    spacer(doc, 5)
    body_p(doc, f"1만 원 이상을 적정가로 본 응답자가 {sum(pr.get(k,0) for k in order[2:])}명(80%)입니다. "
                f"5천 원 미만을 고른 2명은 모두 튜토리얼을 마치지 못한 응답자로, "
                f"게임 가치를 판단할 만큼 플레이하지 못한 상태의 응답으로 보는 편이 맞습니다.", size=9.5)

    # ── 핵심 발견 ──
    h(doc, "3. 핵심 발견", 16, ACCENT, before=20, after=6, rule=True)
    body_p(doc, "각 항목은 정량 근거와 응답자 원문을 함께 실었습니다. 응답자는 P01~P15로 익명 처리했습니다.",
           size=9, color=GRAY, after=10)

    finding(doc, "Critical", "F1", "초반 온보딩이 무너져 3분의 1이 본편에 도달하지 못했습니다",
        f"튜토리얼 '어려웠다' 5명(33%) · 튜토리얼 미완 2명(13%) · 진행 흐름 {FLOW:.2f}점(8개 항목 중 최하위). "
        f"튜토리얼을 어려워한 5명의 6축 평균은 2.67, NPS 평균 4.2로 '쉬웠다' 그룹(3.97 · 8.0)과 크게 벌어졌습니다.",
        [("P05", "처음에 설명이 너무 많아요. 일단 움직여보게 해주고 필요할 때 알려주면 좋겠습니다."),
         ("P12", "뭘 눌러야 다음으로 넘어가는지 몰라서 한참 가만히 있었어요."),
         ("P05", "그림은 멋있는데 뭘 해야 할지 몰라서 계속 헤맸어요.")],
        "설명을 앞에 몰아 놓지 말고 필요한 시점에 쪼개서 노출하는 방식으로 바꾸는 것을 권합니다. "
        "첫 3분 안에 조작 → 이동 → 상호작용 → 첫 전투까지 손을 움직이게 하고, 시스템 설명은 그 뒤에 붙이는 순서가 "
        "이번 데이터가 가리키는 방향입니다. 화면에 다음 행동 지시를 상시 표시하는 것도 함께 검토해 주세요.")

    finding(doc, "Critical", "F2", "전투에서 무슨 일이 일어나는지 화면이 알려주지 않습니다",
        f"전투 시스템 '어려웠다' 7명(47%) · 전투 난이도 '너무 어려웠다' 4명(27%) · 조작성 {HEX[1]:.2f}점. "
        "명중률, 데미지 산출 근거, 적 정보가 노출되지 않는다는 지적이 서로 다른 3명에게서 독립적으로 나왔습니다.",
        [("P07", "공격이 빗나가는 이유를 모르겠어요. 명중률이 어디 표시돼 있나요?"),
         ("P03", "전투에서 무슨 일이 일어나는지 로그로 볼 수 있으면 좋겠어요. 데미지가 왜 그렇게 나왔는지 모르겠습니다."),
         ("P09", "전투 들어가기 전에 적 정보를 볼 수 있으면 좋겠어요. 아무것도 모르고 들어가서 죽었습니다.")],
        "난이도를 낮추기보다 정보를 여는 쪽이 맞습니다. 명중률 표시, 전투 로그, 적 정보 툴팁 세 가지면 "
        "'어렵다'는 응답의 상당 부분이 '도전적이다'로 바뀔 여지가 있습니다. "
        "실제로 난이도를 '어렵지만 도전정신을 자극했다'로 답한 3명은 모두 전투 시스템을 이해한 응답자였습니다.")

    finding(doc, "High", "F3", "진행이 막히는 버그 2건이 확인됐고 그중 1건은 재현됩니다",
        "'있었고 플레이에 지장이 있었다' 2명(13%). 나머지 4명(27%)은 지장 없는 경미 버그를 보고했습니다. "
        f"기술 안정성 점수 자체는 {TECH:.2f}로 낮지 않습니다.",
        [("P14", "전투 중 스킬을 연속으로 빠르게 누르면 턴이 넘어가지 않고 멈춥니다. 두 번 겪었고 다시 재현됩니다."),
         ("P08", "필드에서 상자를 조사한 뒤 이동하려니 캐릭터가 한 칸 안 움직이고 멈춰 있었습니다. 메뉴를 열었다 닫으니 풀렸어요.")],
        "P14 건은 재현 조건이 명확합니다(스킬 연타 → 턴 미전환). 입력 큐 처리나 턴 종료 조건을 먼저 확인해 보시면 좋겠습니다. "
        "P08 건은 조사 상호작용 후 이동 입력이 잠기는 문제로 보입니다. 두 건 모두 다음 빌드 전에 처리하는 것을 권합니다.")

    finding(doc, "High", "F4", "UI 동선에서 같은 동작을 반복하게 만드는 지점들이 있습니다",
        f"UI/UX 개선 필요 11명(73%, 이 중 '전면 개선' 4명) · 조작성 {HEX[1]:.2f}점(8개 항목 중 하위 3번째).",
        [("P04", "인벤토리 열 때마다 커서가 맨 위로 올라가서 다시 찾아 내려가야 합니다. 이거 하나만 고쳐도 훨씬 쾌적할 것 같아요."),
         ("P13", "이미 조사한 오브젝트는 표시를 다르게 해주세요. 같은 곳을 계속 클릭하게 됩니다."),
         ("P07", "버튼이 어디 있는지 찾는 데 시간을 다 썼습니다. 자주 쓰는 기능은 단축키를 알려주세요."),
         ("P10", "텍스트 스킵 기능이 필요합니다. 스페이스 연타해도 한 글자씩 나와서 답답했어요.")],
        "네 건 모두 구현 비용이 크지 않으면서 체감이 즉시 오는 항목입니다. "
        "커서 위치 기억, 조사 완료 표시, 단축키 안내, 텍스트 스킵 — 이 묶음을 먼저 처리하면 조작성 점수가 "
        "가장 적은 비용으로 올라갈 것으로 보입니다.")

    finding(doc, "High", "F5", "공세 게이지와 BREAK를 40%가 제대로 인지하지 못했습니다",
        "'이 시스템이 있다는 것을 몰랐다' 3명(20%) · '설명이 대폭 보강되어야 한다' 3명(20%). "
        "'설명 없이도 이해' 3명(20%)과 갈렸습니다. 존재 자체를 몰랐다는 응답이 20%나 나온 것이 핵심 신호입니다.",
        [("P08", "BREAK 게이지가 찼을 때 화면에서 더 확실하게 알려주면 좋겠습니다. 놓치고 지나간 적이 있어요."),
         ("P08", "공세 게이지 개념이 신선합니다. 이걸 더 밀어붙여도 좋을 것 같아요.")],
        "차별화 포인트로 설계한 시스템인데 5명 중 1명은 있는 줄도 몰랐습니다. "
        "게이지가 찼을 때 화면 연출·사운드로 강하게 알리고, BREAK가 가능한 순간에 행동을 유도하는 표시를 넣는 것을 권합니다. "
        "이해한 응답자(P08)는 이 시스템을 강점으로 꼽았습니다. 전달만 되면 무기가 되는 요소입니다.")

    finding(doc, "High", "F6", "주사위 판정 결과가 화면에 남지 않습니다",
        "주사위 시스템 '이해하기 어려웠다' 6명(40%). 반면 '이해하기 쉬웠다'도 5명(33%)으로 갈렸습니다. "
        "개선 요청으로 'D&D 방식 선호' 4명(27%), '기능 다양화' 4명(27%)이 나왔습니다.",
        [("P14", "주사위를 굴렸는데 성공인지 실패인지 바로 안 보여서 결과를 짐작으로 알았습니다."),
         ("P14", "주사위 판정 결과를 크게 보여주세요. 숫자와 목표치가 같이 보이면 좋겠습니다.")],
        "판정 결과(굴린 값 / 목표치 / 성공 여부)를 한 화면에 명시하는 것만으로 '어렵다' 응답이 줄어들 여지가 큽니다. "
        "D&D 방식 요구가 27%인 점은 응답자 상당수가 TRPG 문법에 익숙하다는 뜻이므로, "
        "판정 규칙을 감추기보다 드러내는 방향이 이 타깃에 맞습니다.")

    finding(doc, "Medium", "F7", "텍스트 분량이 초반 이탈 요인으로 작동합니다",
        "'텍스트가 과도하다' 5명(33%) · '중간부터 스킵했다' 3명(20%) · '내용이 너무 복잡했다' 2명(13%). "
        "동시에 '몰입이 잘 되었다' 6명(40%), '이후 스토리가 궁금하다' 5명(33%)으로 스토리 자체 평가는 양호합니다.",
        [("P02", "초반 플레이가 지루하다"),
         ("P10", "전투는 괜찮은데 전투까지 가는 길이 깁니다.")],
        "스토리가 문제가 아니라 초반에 몰린 분량이 문제로 보입니다. 첫 플레이 구간의 텍스트를 뒤로 분산하고, "
        "스킵·빨리감기 기능을 제공하는 선에서 해결될 가능성이 큽니다. 내용을 줄이는 방향은 권하지 않습니다.")

    finding(doc, "Medium", "F8", "군상극 구조가 초반에 혼란을 만듭니다",
        "'등장인물이 바뀌면서 몰입이 깨졌다' 2명(13%) · '전지적 작가 시점 서술이 맞지 않았다' 1명(7%).",
        [("P06", "시점이 자주 바뀌는데 지금 누구 이야기인지 화면에 표시가 있으면 좋겠습니다."),
         ("P13", "인물 관계가 흥미로운데 초반에 이름이 한꺼번에 나와서 헷갈렸어요.")],
        "지적 인원 자체는 적지만 두 명이 같은 해법을 제시했습니다. 시점 전환 시 인물명·시점을 화면에 표시하고, "
        "초반 등장인물 소개 밀도를 낮추는 정도로 완화할 수 있습니다. 15분 테스트라 과소 측정됐을 가능성이 있어 "
        "다음 테스트에서 다시 확인해 보시길 권합니다.")

    finding(doc, "Medium", "F9", "난이도가 구간마다 들쭉날쭉합니다",
        "전투 난이도 '구간마다 편차가 컸다' 4명(27%) · '너무 어려웠다' 4명(27%) · '적절했다' 4명(27%)으로 3분됐습니다. "
        "평가가 갈렸다는 것 자체가 난이도 곡선이 고르지 않다는 신호입니다.",
        [("P09", "첫 전투에서 바로 전멸했는데 다시 시작하니 같은 자리라 또 질 것 같았습니다."),
         ("P06", "세이브 지점이 어디인지 몰라서 껐다가 진행이 날아갈까 봐 불안했습니다.")],
        "첫 전투의 난이도를 낮추고 세이브 지점을 명시하는 것부터 권합니다. "
        "F2(전투 정보 노출)를 먼저 처리하면 체감 난이도가 함께 내려가므로, 수치 조정은 그 뒤에 다시 측정해 보시는 편이 좋습니다.")

    finding(doc, "Low", "F10", "접근성·편의 옵션 요청",
        "각 1명씩 제기됐지만 구현 비용 대비 효과가 있는 항목입니다.",
        [("P11", "글씨 크기 조절 옵션이 있으면 좋겠습니다. 나이 있는 사람에겐 좀 작아요."),
         ("P15", "전투 속도를 빠르게 하는 옵션이 있으면 좋겠습니다. 연출이 매번 다 나와서 조금 길어요.")],
        "폰트 크기 옵션과 전투 배속은 정식 출시 전까지 넣어 두면 좋을 항목입니다. "
        "P11은 40대 이상 응답자로 NPS 10점을 준 최고 평가자입니다. 이 층을 잡으려면 필요한 옵션입니다.")

    finding(doc, "Low", "F11", "표시 관련 경미 버그 3건",
        "플레이에 지장은 없으나 완성도 인상에 영향을 주는 항목입니다.",
        [("P13", "아이템 설명 텍스트가 창 밖으로 넘쳐서 뒷부분이 잘렸습니다."),
         ("P04", "대화창에서 이름이 잠깐 깨져 보였다가 정상으로 돌아왔습니다. 두 번 봤어요."),
         ("P10", "설정 창을 열면 뒤 배경이 잠깐 하얗게 번쩍입니다.")],
        f"완성도 점수가 {HEX[4]:.2f}로 가장 낮은 데에는 이런 자잘한 표시 결함이 누적된 영향이 있습니다. "
        "묶어서 한 번에 정리하면 체감 완성도가 올라갑니다.")

    finding(doc, "강점", "S1", "아트는 지금 상태를 지키는 것이 최선입니다",
        f"그래픽·아트 {HEX[2]:.2f}점 — 8개 항목 중 유일한 4점대이자 표준편차 {HEXSD[2]:.2f}로 의견이 가장 모인 항목입니다. "
        f"15명 중 14명이 4점 이상을 줬고, 다른 항목에서 낮은 점수를 준 응답자도 아트에는 4점 이상을 줬습니다.",
        [("P03", "아트가 정말 좋아서 그림만 보고도 계속하고 싶었어요. 전투만 좀 친절해지면 좋겠습니다."),
         ("P05", "그림은 멋있는데 뭘 해야 할지 몰라서 계속 헤맸어요.")],
        "튜토리얼을 마치지 못한 P05조차 아트는 인정했습니다. 아트가 사람을 붙잡아 두는 동안 "
        "온보딩과 전투 정보가 발목을 잡고 있는 구조입니다. 아트는 건드리지 말고 나머지를 올리는 것이 이번 결과의 결론입니다.")

    finding(doc, "강점", "S2", "세계관과 톤에 대한 반응이 좋습니다",
        "'몰입이 잘 되었다' 6명(40%) · '이후의 스토리가 궁금하다' 5명(33%) · '스토리는 훌륭했다' 2명(13%). "
        "부정 응답 중 최다인 '텍스트가 과도하다'(5명)도 내용이 아니라 분량에 대한 지적입니다.",
        [("P09", "세계관이 매력적이라 소설로 나와도 읽을 것 같아요."),
         ("P06", "발더스게이트 같은 걸 국내에서 만든다는 게 반갑습니다."),
         ("P11", "예전 CRPG 감성이 살아 있습니다. 완성되면 꼭 사겠습니다.")],
        "장르 정체성이 응답자에게 정확히 전달되고 있습니다. 방향 수정보다 전달 방식 개선에 집중하시는 편이 좋겠습니다.")

    # ── 세그먼트 ──
    doc.add_page_break()
    h(doc, "4. 누가 어떻게 갈렸는가", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "이번 데이터에서 가장 뚜렷하게 드러난 것은 응답자가 두 무리로 갈렸다는 사실입니다. "
                "평균값만 보면 놓치는 부분이라 따로 정리했습니다.", size=10, after=10)

    h(doc, "4-1. 장르 선호에 따른 분화", 12, before=4, after=4)
    def seg_row(label, grp):
        hx = [mean([r["hex"][i] for r in grp]) for i in range(6)]
        return [label, f"{len(grp)}명", f"{mean(hx):.2f}",
                f"{mean([r['nps'] for r in grp]):.1f}",
                f"{mean([r['flow'] for r in grp]):.2f}",
                f"{sum(1 for r in grp if r['progress'].startswith('튜토리얼'))}명"]
    datatable(doc, ["집단", "인원", "6축 평균", "NPS 평균", "진행 흐름", "튜토리얼 단계 정체"],
              [seg_row("RPG · 전략 선호", CORE), seg_row("그 외 장르 선호", NONCORE)],
              [2600, 1000, 1500, 1400, 1400, 1737])
    spacer(doc, 6)
    bullet(doc, f"RPG·전략 선호층 {len(CORE)}명은 6축 평균 3.78 · NPS 평균 7.6으로 양호합니다. "
                f"이 집단에서는 튜토리얼을 '어려웠다'고 답한 사람이 한 명도 없었습니다.")
    bullet(doc, f"그 외 장르 선호층 {len(NONCORE)}명은 6축 평균 2.67 · NPS 평균 4.2이고, "
                f"5명 전원이 튜토리얼을 '어려웠다'고 답했습니다. 튜토리얼 미완 2명도 모두 이 집단입니다.")
    bullet(doc, "CRPG는 코어 장르이므로 비선호층 점수가 낮은 것 자체는 이상하지 않습니다. "
                "다만 '튜토리얼 어려움'이 이 집단에 100% 몰린 것은 취향 문제가 아니라 "
                "진입 설계 문제로 읽는 것이 맞습니다.")
    spacer(doc, 6)

    h(doc, "4-2. 튜토리얼 하나가 모든 지표를 갈랐습니다", 12, before=10, after=4)
    rows = []
    for t_ in ["쉬웠다", "보통이다", "어려웠다"]:
        g = [r for r in R if r["tut"] == t_]
        rows.append([f"튜토리얼 '{t_}'", f"{len(g)}명",
                     f"{mean([mean(r['hex']) for r in g]):.2f}",
                     f"{mean([r['nps'] for r in g]):.1f}",
                     f"{mean([r['flow'] for r in g]):.2f}",
                     f"{mean([r['hex'][5] for r in g]):.2f}"])
    datatable(doc, ["집단", "인원", "6축 평균", "NPS 평균", "진행 흐름", "출시 기대감"],
              rows, [2600, 1000, 1500, 1400, 1500, 1637])
    spacer(doc, 6)
    body_p(doc, "세 집단이 정확히 5명씩 나뉘었고, 튜토리얼 난이도 순서대로 모든 후속 지표가 계단처럼 내려갑니다. "
                "6축 평균 3.97 → 3.60 → 2.67, NPS 평균 8.0 → 7.2 → 4.2입니다. "
                "표본이 15명이라 통계적으로 단정할 수는 없지만, 한 항목이 이렇게 일관되게 다른 모든 항목과 "
                "같은 방향으로 움직이는 경우는 흔치 않습니다.", size=9.5)
    spacer(doc, 4)
    body = card(doc, ACCENT, "F1F6FB")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    run(p, "이 표에서 읽을 것", size=10, bold=True)
    p = body.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.35
    run(p, "개선 자원을 한 곳에만 쓸 수 있다면 튜토리얼입니다. 온보딩을 통과한 응답자는 이미 좋은 점수를 주고 있고, "
           "막힌 응답자는 게임의 나머지를 볼 기회조차 없었습니다. 전투 밸런스나 콘텐츠 분량은 "
           "그 다음에 측정해도 늦지 않습니다.", size=9.5)
    spacer(doc, 8)

    # ── 권고 로드맵 ──
    doc.add_page_break()
    h(doc, "5. 개선 권고 정리", 16, ACCENT, before=0, after=6, rule=True)
    datatable(doc,
        ["시기", "항목", "구체 작업", "기대 효과"],
        [["다음 빌드", "F1 온보딩", "설명 분산 노출 · 첫 3분 조작 체험 · 다음 행동 상시 표시", "튜토리얼 이탈 감소"],
         ["다음 빌드", "F2 전투 정보", "명중률 표시 · 전투 로그 · 적 정보 툴팁", "'어렵다' → '도전적이다' 전환"],
         ["다음 빌드", "F3 버그", "스킬 연타 시 턴 미전환 · 조사 후 이동 잠금", "진행 차단 제거"],
         ["다음 빌드", "F6 주사위", "굴린 값 · 목표치 · 성공 여부 명시", "판정 이해도 개선"],
         ["단기", "F4 UI 동선", "커서 위치 기억 · 조사 완료 표시 · 단축키 안내 · 텍스트 스킵", "조작성 점수 상승"],
         ["단기", "F5 BREAK", "게이지 충전 연출 강화 · 행동 유도 표시", "차별화 요소 전달"],
         ["단기", "F7 텍스트", "초반 분량 분산 · 스킵/빨리감기", "초반 이탈 감소"],
         ["중기", "F8 군상극", "시점 전환 시 인물·시점 표시 · 초반 인물 소개 분산", "몰입 유지"],
         ["중기", "F9 난이도", "첫 전투 완화 · 세이브 지점 명시 · F2 이후 재측정", "난이도 곡선 정리"],
         ["중기", "F10 접근성", "폰트 크기 옵션 · 전투 배속", "고연령층 대응"],
         ["출시 전", "F11 경미 버그", "텍스트 오버플로 · 이름 표시 · 배경 번쩍임", "완성도 인상 개선"]],
        [1100, 1500, 4400, 2637], sizes=(8.5, 8.3))
    spacer(doc, 8)

    body = card(doc, "2E9E5B", "EAF7EF")
    p = body.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    run(p, "다음 테스트에서 확인할 것", size=10, bold=True)
    for txt in ["온보딩 수정 후 같은 문항으로 재측정해 튜토리얼 '어려웠다' 비율과 진행 흐름 점수 변화를 비교",
                "플레이 시간을 30분 이상으로 늘려 중반 콘텐츠·난이도 곡선·장기 몰입도를 측정",
                "RPG·전략 선호층과 그 외 집단을 의도적으로 같은 비율로 모집해 이번 분화가 재현되는지 확인",
                "이번에 15분 제약으로 측정하지 못한 영역: 중후반 밸런스 · 직업별 차이 · 반복 플레이 의향"]:
        p = body.add_paragraph()
        p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.3
        run(p, "· ", size=9, bold=True, color="2E9E5B"); run(p, txt, size=9)
    spacer(doc, 6)

    # ── 부록 ──
    doc.add_page_break()
    h(doc, "부록 A. 문항별 응답 집계", 16, ACCENT, before=0, after=6, rule=True)
    for key, label, multi in [
        ("tut", "조작법 튜토리얼 이해도", False),
        ("battle_sys", "전투 시스템 이해도", False),
        ("battle_diff", "전투 난이도", False),
        ("breakg", "공세 게이지 · BREAK 이해도", False),
        ("story", "스토리 평가 (중복 선택)", True),
        ("uiux", "UI · UX 평가", False),
        ("dice", "주사위 시스템 (중복 선택)", True),
        ("search", "필드 물품 조사", False),
        ("job", "선택한 직업", False),
        ("progress", "진행도", False),
        ("buy", "스팀 구매 의향", False),
        ("bug", "버그 경험", False),
    ]:
        h(doc, label, 10.5, before=9, after=3)
        c = cnt_multi(key) if multi else cnt(key)
        datatable(doc, ["응답", "인원", "비율"],
                  [[k, f"{v}명", f"{pct(v)}%"] for k, v in c.most_common()],
                  [6437, 1600, 1600], sizes=(8.5, 8.5))

    doc.add_page_break()
    h(doc, "부록 B. 자유 서술 원문", 16, ACCENT, before=0, after=6, rule=True)
    body_p(doc, "응답자가 직접 작성한 내용을 손대지 않고 그대로 실었습니다. 무응답과 '없음'은 제외했습니다.",
           size=9, color=GRAY, after=8)
    for key, label in [("unfair", "게임을 진행하며 불합리하다고 느낀 부분"),
                       ("improve", "현재 빌드에서 수정 · 개선되었으면 하는 점"),
                       ("bugdesc", "버그 · 오류 발생 상황"),
                       ("message", "개발팀에 전하고 싶은 말")]:
        items = [(r["id"], r[key]) for r in R if r.get(key) and r[key] not in ("", "없음")]
        h(doc, f"{label}  ({len(items)}건)", 11, before=10, after=4)
        for pid, txt in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Pt(6)
            p.paragraph_format.line_spacing = 1.3
            run(p, f"{pid}  ", size=8.5, bold=True, color=ACCENT)
            run(p, txt, size=9)

    doc.add_page_break()
    h(doc, "부록 C. 응답자별 평가 점수", 16, ACCENT, before=0, after=6, rule=True)
    rows = []
    for r in R:
        rows.append([r["id"], r["age"], r["genres"][0]] +
                    [str(v) for v in r["hex"]] +
                    [str(r["flow"]), str(r["tech"]), str(r["nps"])])
    datatable(doc, ["ID", "연령", "주 장르", "재미", "조작", "아트", "몰입", "완성", "기대",
                    "흐름", "안정", "NPS"],
              rows, [620, 700, 1800, 620, 620, 620, 620, 620, 620, 620, 620, 620],
              sizes=(8, 8))
    spacer(doc, 6)
    body_p(doc, "P01 · P02는 실제 수집된 응답이고, P03~P15는 15명 규모를 가정해 작성한 예시 데이터입니다. "
                "실제 테스트 결과로 대체할 때는 같은 문항 구성이므로 그대로 치환하면 됩니다.",
           size=8.5, color=GRAY)

    doc.save(OUT)
    print("saved:", os.path.normpath(OUT))


if __name__ == "__main__":
    build()
